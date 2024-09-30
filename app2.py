import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from io import BytesIO
import zipfile
from docx2pdf import convert  # Aggiungi docx2pdf per convertire in PDF
import os
import pythoncom
import win32com.client as win32

# Funzione per convertire un documento Word in PDF
def convert_to_pdf(word_file, output_pdf_path):
    # Inizializza COM
    pythoncom.CoInitialize()

    try:
        with open("temp_decreto.docx", "wb") as temp_word_file:
            temp_word_file.write(word_file.getvalue())

        # Crea un'istanza di Word
        word = win32.Dispatch("Word.Application")
        doc = word.Documents.Open(os.path.abspath("temp_decreto.docx"))

        # Salva il documento come PDF
        doc.SaveAs(os.path.abspath(output_pdf_path), FileFormat=17)  # 17 è il formato PDF in Word
        doc.Close()

        # Chiudi l'applicazione Word
        word.Quit()

        # Rimuovi il file temporaneo
        os.remove("temp_decreto.docx")
    
    finally:
        # Deinizializza COM
        pythoncom.CoUninitialize()

# Funzione per caricare i file Excel o CSV
def carica_file():
    uploaded_file1 = st.file_uploader("Carica il file Anagrafiche (Excel o CSV)", type=["xlsx", "csv"])
    uploaded_file2 = st.file_uploader("Carica il file Fatture (Excel o CSV)", type=["xlsx", "csv"])
    uploaded_file3 = st.file_uploader("Carica il file Pratiche (Excel o CSV)", type=["xlsx", "csv"])
    
    def leggi_file(file):
        if file is not None:
            try:
                if file.name.endswith('.xlsx'):
                    return pd.read_excel(file)
                elif file.name.endswith('.csv'):
                    return pd.read_csv(file, sep=';', on_bad_lines='skip')
            except Exception as e:
                st.error(f"Errore nel caricamento del file {file.name}: {e}")
        return None

    df1 = leggi_file(uploaded_file1)  # File Anagrafiche
    df2 = leggi_file(uploaded_file2)  # File Fatture
    df3 = leggi_file(uploaded_file3)  # File Pratiche
    
    if df1 is None or df2 is None or df3 is None:
        st.warning("Uno o più file non sono stati caricati correttamente.")
        return None, None, None
    
    return df1, df2, df3

# Funzione per normalizzare i nomi delle colonne (convertirli in minuscolo e rimuovere caratteri speciali)
def normalizza_colonne(df):
    df.columns = df.columns.str.strip().str.lower().str.replace(' ', '_').str.replace('.', '').str.replace("'", "")
    return df

# Funzione per creare una tabella delle fatture
def crea_tabella_fatture(doc, df_combinato):
    # Aggiungi una tabella con 7 colonne
    table = doc.add_table(rows=1, cols=7)

    # Aggiungi intestazioni di colonna
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Data reg.'
    hdr_cells[1].text = 'Scad. netto'
    hdr_cells[2].text = 'N. documento'
    hdr_cells[3].text = 'Importo totale'
    hdr_cells[4].text = 'Importo pagato totale'
    hdr_cells[5].text = 'Residuo ad oggi'
    hdr_cells[6].text = 'Punto fornitura (POD)'

    # Popola la tabella con le fatture
    for index, row in df_combinato.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row.get('data_reg', ''))  # Data reg.
        cells[1].text = str(row.get('scad_netto', ''))  # Scad. netto
        cells[2].text = str(row.get('n_documento', ''))  # N. documento
        cells[3].text = str(row.get('importo_totale', ''))  # Importo totale
        cells[4].text = str(row.get('importo_pagato_totale', ''))  # Importo pagato totale
        cells[5].text = str(row.get('residuo_ad_oggi', ''))  # Residuo ad oggi
        cells[6].text = str(row.get('pod', ''))  # Punto fornitura (POD)

# Funzione per generare il documento Word con sostituzioni e aggiunta della tabella
def genera_documento_word(dati, df_combinato, template_path="decreto.docx"):
    st.write("Dati per la sostituzione:", dati)
    doc = Document(template_path)

    # Sostituzione dei segnaposto nel documento
    placeholders = {
        "{ragione_sociale}": str(dati.get('ragione_sociale', '')).replace('nan', ''),
        "{codice_fiscale}": str(dati.get('codice_fiscale', '')).replace('nan', ''),
        "{partita_va}": str(dati.get('partita_iva', '')).replace('nan', ''),
        "{comune_residenza}": str(dati.get('comune_residenza', '')).replace('nan', ''),
        "{cap_residenza}": str(int(dati.get('cap_residenza', 0))),
        "{indirizzo_residenza}": str(dati.get('indirizzo_residenza', '')).replace('nan', ''),
        "{settore_contabile}": str(dati.get('settore_contabile', '')).replace('nan', ''),
        "{codice_commerciale}": str(dati.get('codice_commerciale', '')).replace('nan', ''),
        "{codice_soggetto}": str(int(dati.get('codice_soggetto', 0))),
        "{comune_fornitura}": str(dati.get('comune_fornitura', '')).replace('nan', ''),
        "{provincia_fornitura}": str(dati.get('provincia_fornitura', '')).replace('nan', ''),
        "{indirizzo_fornitura}": str(dati.get('indirizzo_fornitura', '')).replace('nan', ''),
        "{pod}": str(dati.get('pod', '')).replace('nan', ''),
        "{residuo_ad_oggi}": str(dati.get('residuo_ad_oggi', '')).replace('nan', ''),
    }
    
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # Cerca il segnaposto per la tabella
    for paragraph in doc.paragraphs:
        if "{tabella_fatture}" in paragraph.text:
            # Cancella il paragrafo del segnaposto
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None
            
            # Aggiungi la tabella al posto del segnaposto
            crea_tabella_fatture(doc, df_combinato)
            break

    # Salva il documento in un buffer BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Funzione principale
def main():
    st.title("Generatore di Documenti Word e PDF da Excel o CSV")
    df1, df2, df3 = carica_file()

    if df1 is not None and df2 is not None and df3 is not None:
        st.write("File caricati con successo!")

        try:
            # Normalizza i nomi delle colonne a minuscolo
            df1 = normalizza_colonne(df1)
            df2 = normalizza_colonne(df2)
            df3 = normalizza_colonne(df3)

            # Controlla le colonne disponibili
            st.write("Colonne del file Anagrafiche (df1):", df1.columns)
            st.write("Colonne del file Fatture (df2):", df2.columns)
            st.write("Colonne del file Pratiche (df3):", df3.columns)

            # Rinominazione delle colonne per fare il merge
            df1 = df1.rename(columns={"codice_soggetto": "codice_soggetto"})
            df2 = df2.rename(columns={"bpartner": "codice_soggetto"})
            df3 = df3.rename(columns={"soggetto": "codice_soggetto"})
        except KeyError as e:
            st.error(f"Errore durante la rinominazione delle colonne: {e}")
            return

        if "codice_soggetto" not in df1.columns or "codice_soggetto" not in df2.columns or "codice_soggetto" not in df3.columns:
            st.error("Una o più colonne 'codice_soggetto' non sono state trovate.")
            return

        try:
            df_combinato = pd.merge(df1, df2, on='codice_soggetto', how='left')  # Unione tra df1 e df2
            df_combinato = pd.merge(df_combinato, df3, on='codice_soggetto', how='left')  # Unione con df3
        except KeyError as e:
            st.error(f"Errore durante l'unione dei file: {e}")
            return

        codici_soggetto = df_combinato['codice_soggetto'].unique()
        codici_selezionati = st.multiselect("Seleziona i codici soggetto per generare i documenti:", codici_soggetto)

        if codici_selezionati:
            zip_buffer = BytesIO()

            with zipfile.ZipFile(zip_buffer, "w") as zip_file:
                for codice in codici_selezionati:
                    # Filtra i dati per il codice soggetto selezionato
                    dati_filtro = df_combinato[df_combinato['codice_soggetto'] == codice].iloc[0]
                    
                    # Filtra le fatture corrispondenti a questo codice soggetto
                    df_fatture = df_combinato[df_combinato['codice_soggetto'] == codice]

                    # Genera il documento Word con la tabella delle fatture
                    doc_buffer = genera_documento_word(dati_filtro, df_fatture, template_path="decreto.docx")
                    
                    # Nome del file personalizzato: 'Ragione_Sociale' + 'Codice_Soggetto' + 'Numero_Affido'
                    nome_file = f"{int(dati_filtro['codice_soggetto'])}"
                    
                    # Aggiungi il documento Word allo zip
                    zip_file.writestr(f"{nome_file}.docx", doc_buffer.getvalue())
                    
                    # Genera anche il PDF e aggiungilo allo zip
                    pdf_path = f"{nome_file}.pdf"
                    convert_to_pdf(doc_buffer, pdf_path)
                    with open(pdf_path, "rb") as pdf_file:
                        zip_file.writestr(f"{nome_file}.pdf", pdf_file.read())
                    
                    # Rimuovi il PDF generato
                    os.remove(pdf_path)
            
            zip_buffer.seek(0)
            st.download_button(
                label="Scarica tutti i documenti generati (ZIP)",
                data=zip_buffer,
                file_name="documenti_generati.zip",
                mime="application/zip"
            )
    else:
        st.warning("Per favore, carica tutti e tre i file (Excel o CSV).")

if __name__ == "__main__":
    main()
